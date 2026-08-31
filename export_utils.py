"""
============================================================
EduMindAI Enterprise v3.0
Export Utilities (PDF & DOCX)
============================================================
"""

import io
from docx import Document
from fpdf import FPDF


class ExportManager:

    @staticmethod
    def to_docx(messages):
        """Chat tarixini Word (.docx) fayliga o'girish"""
        doc = Document()
        doc.add_heading("EduMindAI - Chat Hisoboti", level=1)

        for msg in messages:
            role = "Foydalanuvchi" if msg["role"] == "user" else "EduMindAI"
            content = msg["content"]
            doc.add_paragraph(f"{role}:", style="Heading 2")
            doc.add_paragraph(content)
            doc.add_paragraph("-" * 30)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def to_pdf(messages):
        """Chat tarixini PDF fayliga o'girish"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)

        pdf.cell(200, 10, txt="EduMindAI - Chat History", ln=1, align="C")
        pdf.ln(10)

        for msg in messages:
            role = "User" if msg["role"] == "user" else "EduMindAI"
            content = msg["content"].encode('latin-1', 'replace').decode('latin-1')

            pdf.set_font("Helvetica", 'B', 11)
            pdf.cell(0, 8, txt=f"{role}:", ln=1)

            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, txt=content)
            pdf.ln(4)

        return pdf.output()


exporter = ExportManager()
